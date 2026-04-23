import os
import re
import sys
import json
import base64
import subprocess
import threading
import webbrowser
from datetime import datetime
from urllib.parse import quote, quote_plus, urlparse
from tkinter import Tk, Label, Entry, Button, filedialog, StringVar, messagebox, simpledialog, Toplevel, Text
from tkinter import colorchooser
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO
import requests

# ---------------- Paths for executable-friendly ----------------
if getattr(sys, 'frozen', False):
    BASE_DIR = os.path.dirname(sys.executable)
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

LOGO_FOLDER = os.path.join(BASE_DIR, "NewsLogos")
CUSTOM_LOGO_NAME = "custom"
NOTE_FOLDER = os.path.join(BASE_DIR, "Note")
os.makedirs(NOTE_FOLDER, exist_ok=True)
MISSING_LOG_FILE = os.path.join(NOTE_FOLDER, "missing_logos.txt")
SETTINGS_FILE = os.path.join(NOTE_FOLDER, "settings.json")
SUPPORTED_LOGO_EXTENSIONS = (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tif", ".tiff", ".ico", ".jfif")
APP_VERSION = "2.1.0"
UPDATE_INFO_URL = "https://raw.githubusercontent.com/deepndense-sketch/PrintNews/main/version.json"
GITHUB_REPO_OWNER = "deepndense-sketch"
GITHUB_REPO_NAME = "PrintNews"
GITHUB_BRANCH = "main"
GITHUB_LOGO_API_URL = f"https://api.github.com/repos/{GITHUB_REPO_OWNER}/{GITHUB_REPO_NAME}/contents/NewsLogos?ref={GITHUB_BRANCH}"
GITHUB_CONTENTS_API_URL = f"https://api.github.com/repos/{GITHUB_REPO_OWNER}/{GITHUB_REPO_NAME}/contents"
GITHUB_TOKEN_SETTINGS_KEY = "github_token"
REQUEST_HEADERS = {"User-Agent": "PrintNews"}
CROSS_CHECK_SETTINGS_KEY = "cross_check_news_with_link"
HIGHLIGHT_COLOR_SETTINGS_KEY = "highlight_color"
HIGHLIGHT_OPACITY_SETTINGS_KEY = "highlight_opacity"
DEFAULT_HIGHLIGHT_COLOR = "#fff176"
DEFAULT_HIGHLIGHT_OPACITY = 80
EXPORT_PREFIX = "SourceListPR_"

# ---------------- Config ----------------
PADDING_TOP = 20
PADDING_BOTTOM = 20
MARGIN = 20
FONT_SIZE_HEADLINE = 48
FONT_SIZE_DATE = 24
FONT_SIZE_SOURCE = 20
LOGO_HEIGHT = 60
MAX_FILENAME_WORDS = 8
LINE_SPACING = 10
RIGHT_MARGIN = 25
MIN_WIDTH = 800
MAX_WIDTH = 1500
SUB_HEAD_COLOR = "#003366"  # dark blue for sub headline
GAP_BETWEEN_SEGMENTS = 20  # extra gap for // segments
HIGHLIGHT_PADDING = 5

MONTH_MAP = {
    "jan": 1, "january": 1, "feb": 2, "february": 2, "mar": 3, "march": 3,
    "apr": 4, "april": 4, "may": 5, "jun": 6, "june": 6, "jul": 7, "july": 7,
    "aug": 8, "august": 8, "sep": 9, "sept": 9, "september": 9, "oct": 10,
    "october": 10, "nov": 11, "november": 11, "dec": 12, "december": 12,
}


def version_parts(version):
    parts = []
    for piece in re.findall(r"\d+", version or ""):
        parts.append(int(piece))
    return parts or [0]


def is_newer_version(remote_version, current_version):
    remote_parts = version_parts(remote_version)
    current_parts = version_parts(current_version)
    length = max(len(remote_parts), len(current_parts))
    remote_parts += [0] * (length - len(remote_parts))
    current_parts += [0] * (length - len(current_parts))
    return remote_parts > current_parts


def fetch_update_info():
    response = requests.get(UPDATE_INFO_URL, timeout=6)
    response.raise_for_status()
    info = json.loads(response.content.decode("utf-8-sig"))
    latest_version = str(info.get("version", "")).strip()
    return {
        "version": latest_version,
        "notes": str(info.get("notes", "")).strip(),
        "download_url": str(info.get("download_url", "")).strip(),
        "is_newer": bool(latest_version and is_newer_version(latest_version, APP_VERSION)),
    }


def check_for_updates(show_current=False):
    try:
        info = fetch_update_info()
        latest_version = info["version"]
        if info["is_newer"]:
            message = f"Update available to version {latest_version}.\n\nCurrent version: {APP_VERSION}\nLatest version: {latest_version}"
            if info["notes"]:
                message += f"\n\nWhat is updated:\n{info['notes']}"
            if info["download_url"]:
                message += f"\n\nDownload/update link:\n{info['download_url']}"
            messagebox.showinfo("Update Available", message)
        elif show_current:
            latest_label = latest_version or "unknown"
            messagebox.showinfo("No Update", f"No update available.\n\nCurrent version: {APP_VERSION}\nLatest version: {latest_label}")
    except Exception:
        if show_current:
            messagebox.showwarning("Update Check Failed", "Could not check for updates right now.")


update_info = None


def update_button_from_info(info=None, error=None):
    global update_info
    if error:
        update_button.config(text="Update check failed")
        return
    update_info = info
    latest_version = info.get("version") or APP_VERSION
    if info.get("is_newer"):
        update_button.config(text=f"Update to {latest_version}")
    else:
        update_button.config(text=f"Latest version {latest_version}")


def run_update_button_check():
    try:
        info = fetch_update_info()
        root.after(0, update_button_from_info, info, None)
    except Exception as e:
        root.after(0, update_button_from_info, None, e)


def start_update_button_check():
    update_button.config(text="Checking update...")
    threading.Thread(target=run_update_button_check, daemon=True).start()


def install_update(info):
    download_url = info.get("download_url")
    latest_version = info.get("version")
    if not download_url:
        root.after(0, messagebox.showwarning, "Update Failed", "The update file link is missing.")
        return
    if not getattr(sys, "frozen", False):
        root.after(0, messagebox.showinfo, "Update Available", f"Update available to version {latest_version}.\n\nDownload link:\n{download_url}")
        return

    try:
        root.after(0, update_button.config, {"text": f"Downloading {latest_version}..."})
        response = requests.get(download_url, timeout=60)
        response.raise_for_status()
        exe_path = sys.executable
        new_exe_path = exe_path + ".new"
        updater_path = os.path.join(BASE_DIR, "apply_update.bat")
        with open(new_exe_path, "wb") as f:
            f.write(response.content)
        bat = f"""@echo off
timeout /t 2 /nobreak > nul
move /y "{new_exe_path}" "{exe_path}" > nul
start "" "{exe_path}"
del "%~f0"
"""
        with open(updater_path, "w", encoding="ascii") as f:
            f.write(bat)
        subprocess.Popen(["cmd", "/c", updater_path], cwd=BASE_DIR)
        os._exit(0)
    except Exception as e:
        root.after(0, update_button.config, {"text": f"Update to {latest_version}"})
        root.after(0, messagebox.showwarning, "Update Failed", f"Could not update to version {latest_version}.\n\n{e}")


def get_github_token():
    return (
        os.environ.get("PRINTNEWS_GITHUB_TOKEN")
        or os.environ.get("GITHUB_TOKEN")
        or os.environ.get("GH_TOKEN")
        or str(load_settings().get(GITHUB_TOKEN_SETTINGS_KEY, "")).strip()
    )


def github_headers(token=None):
    headers = {
        "Accept": "application/vnd.github+json",
        "User-Agent": "PrintNews",
        "X-GitHub-Api-Version": "2022-11-28",
    }
    if token:
        headers["Authorization"] = f"Bearer {token}"
    return headers


def upload_logo_to_github(filename, token):
    path_for_api = quote(f"NewsLogos/{filename}", safe="/")
    api_url = f"{GITHUB_CONTENTS_API_URL}/{path_for_api}"
    local_path = os.path.join(LOGO_FOLDER, filename)
    with open(local_path, "rb") as f:
        encoded_content = base64.b64encode(f.read()).decode("ascii")

    payload = {
        "message": f"Add logo {filename}",
        "content": encoded_content,
        "branch": GITHUB_BRANCH,
    }
    response = requests.put(api_url, headers=github_headers(token), json=payload, timeout=20)
    response.raise_for_status()


def sync_logos_with_github():
    try:
        os.makedirs(LOGO_FOLDER, exist_ok=True)
        token = get_github_token()
        local_files = [
            name for name in os.listdir(LOGO_FOLDER)
            if name.lower().endswith(SUPPORTED_LOGO_EXTENSIONS)
            and os.path.isfile(os.path.join(LOGO_FOLDER, name))
        ]
        existing = {name.lower() for name in local_files}
        response = requests.get(GITHUB_LOGO_API_URL, headers=github_headers(token), timeout=10)
        response.raise_for_status()
        remote_files = response.json()
        remote_names = {
            item.get("name", "").lower()
            for item in remote_files
            if item.get("name", "").lower().endswith(SUPPORTED_LOGO_EXTENSIONS)
        }
        downloaded = []
        uploaded = []
        skipped = 0

        for item in remote_files:
            filename = item.get("name", "")
            if not filename.lower().endswith(SUPPORTED_LOGO_EXTENSIONS):
                continue
            if filename.lower() in existing:
                skipped += 1
                continue
            download_url = item.get("download_url")
            if not download_url:
                continue

            logo_response = requests.get(download_url, headers={"User-Agent": "PrintNews"}, timeout=20)
            logo_response.raise_for_status()
            save_path = os.path.join(LOGO_FOLDER, filename)
            if os.path.exists(save_path):
                skipped += 1
                existing.add(filename.lower())
                continue
            with open(save_path, "wb") as f:
                f.write(logo_response.content)
            existing.add(filename.lower())
            downloaded.append(filename)

        upload_skipped = 0
        upload_errors = []
        if token:
            for filename in local_files:
                if filename.lower() in remote_names:
                    upload_skipped += 1
                    continue
                try:
                    upload_logo_to_github(filename, token)
                    remote_names.add(filename.lower())
                    uploaded.append(filename)
                except requests.HTTPError as e:
                    response = getattr(e, "response", None)
                    status_code = response.status_code if response is not None else "unknown"
                    upload_errors.append(f"{filename}: GitHub returned {status_code}")
                except Exception as e:
                    upload_errors.append(f"{filename}: {e}")
        else:
            upload_errors.append("Click Set GitHub Token once on this computer to upload local logos to GitHub.")

        upload_error = None
        if upload_errors:
            upload_error = "\n".join(upload_errors[:10])
            if len(upload_errors) > 10:
                upload_error += f"\n...and {len(upload_errors) - 10} more"

        return downloaded, uploaded, skipped, upload_skipped, upload_error, None
    except Exception as e:
        return [], [], 0, 0, None, e


def show_logo_sync_result(downloaded, uploaded, skipped, upload_skipped, upload_error, error):
    if error:
        messagebox.showwarning("Logo Sync Failed", f"Could not sync logos from GitHub.\n\n{error}")
        return

    message = (
        "Logo sync complete.\n\n"
        f"Downloaded from GitHub: {len(downloaded)}\n"
        f"Uploaded to GitHub: {len(uploaded)}\n"
        f"Already had locally: {skipped}\n"
        f"Already on GitHub: {upload_skipped}"
    )
    if downloaded:
        message += "\n\nDownloaded logos:\n" + "\n".join(downloaded[:30])
        if len(downloaded) > 30:
            message += f"\n...and {len(downloaded) - 30} more"
    if uploaded:
        message += "\n\nUploaded logos:\n" + "\n".join(uploaded[:30])
        if len(uploaded) > 30:
            message += f"\n...and {len(uploaded) - 30} more"
    if upload_error:
        message += f"\n\nUpload skipped:\n{upload_error}"
    messagebox.showinfo("Logo Sync", message)


def run_logo_sync_thread():
    downloaded, uploaded, skipped, upload_skipped, upload_error, error = sync_logos_with_github()
    root.after(0, show_logo_sync_result, downloaded, uploaded, skipped, upload_skipped, upload_error, error)


def start_logo_sync():
    threading.Thread(target=run_logo_sync_thread, daemon=True).start()


def check_updates_clicked():
    if update_info and update_info.get("is_newer"):
        threading.Thread(target=install_update, args=(update_info,), daemon=True).start()
    else:
        start_update_button_check()
# ---------------- GUI ----------------
file_path = None
output_path = None
highlight_bold_text = True
highlight_color = DEFAULT_HIGHLIGHT_COLOR
highlight_opacity = DEFAULT_HIGHLIGHT_OPACITY
action_mode = ""


def load_settings():
    try:
        with open(SETTINGS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def save_settings(word_file, output_folder, bold_highlight_color, bold_highlight_opacity):
    try:
        os.makedirs(NOTE_FOLDER, exist_ok=True)
        settings_data = load_settings()
        settings_data["last_word_file"] = word_file
        settings_data["last_output_folder"] = output_folder
        settings_data[HIGHLIGHT_COLOR_SETTINGS_KEY] = (bold_highlight_color or DEFAULT_HIGHLIGHT_COLOR).strip()
        settings_data[HIGHLIGHT_OPACITY_SETTINGS_KEY] = int(bold_highlight_opacity)
        with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
            json.dump(settings_data, f, indent=2)
    except Exception as e:
        print("Could not save settings:", e)


def save_github_token(token):
    try:
        os.makedirs(NOTE_FOLDER, exist_ok=True)
        settings_data = load_settings()
        settings_data[GITHUB_TOKEN_SETTINGS_KEY] = token
        with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
            json.dump(settings_data, f, indent=2)
        return True
    except Exception as e:
        messagebox.showwarning("Token Not Saved", f"Could not save GitHub token.\n\n{e}")
        return False


def set_github_token():
    token = simpledialog.askstring(
        "Set GitHub Token",
        "Paste a GitHub token with Contents read/write permission for this repo.\n\n"
        "It will be saved on this computer only.",
        show="*",
        parent=root,
    )
    if token is None:
        return
    token = token.strip()
    if not token:
        messagebox.showwarning("Missing Token", "No token was entered.")
        return
    if save_github_token(token):
        messagebox.showinfo("GitHub Token Saved", "Done. Sync Logos can now upload from this computer.")


def save_render_settings():
    try:
        os.makedirs(NOTE_FOLDER, exist_ok=True)
        settings_data = load_settings()
        settings_data[HIGHLIGHT_COLOR_SETTINGS_KEY] = highlight_color_var.get().strip() or DEFAULT_HIGHLIGHT_COLOR
        try:
            settings_data[HIGHLIGHT_OPACITY_SETTINGS_KEY] = int(highlight_opacity_var.get().strip() or DEFAULT_HIGHLIGHT_OPACITY)
        except ValueError:
            settings_data[HIGHLIGHT_OPACITY_SETTINGS_KEY] = DEFAULT_HIGHLIGHT_OPACITY
        with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
            json.dump(settings_data, f, indent=2)
    except Exception as e:
        messagebox.showwarning("Settings Not Saved", f"Could not save render settings.\n\n{e}")


def attach_tooltip(widget, text):
    tooltip = {"window": None}

    def show_tooltip(_event=None):
        if tooltip["window"] or not text:
            return
        x = widget.winfo_rootx() + 10
        y = widget.winfo_rooty() + widget.winfo_height() + 4
        win = Toplevel(widget)
        win.wm_overrideredirect(True)
        win.wm_geometry(f"+{x}+{y}")
        Label(win, text=text, bg="#fffbe6", fg="#333333", relief="solid", bd=1, padx=8, pady=4, wraplength=500, justify="left").pack()
        tooltip["window"] = win

    def hide_tooltip(_event=None):
        if tooltip["window"] is not None:
            tooltip["window"].destroy()
            tooltip["window"] = None

    widget.bind("<Enter>", show_tooltip)
    widget.bind("<Leave>", hide_tooltip)


def choose_highlight_color():
    selected = colorchooser.askcolor(color=highlight_color_var.get().strip() or DEFAULT_HIGHLIGHT_COLOR, parent=root)
    if selected and selected[1]:
        highlight_color_var.set(selected[1])
        save_render_settings()


def source_display_name(source):
    cleaned = (source or "Unknown").strip()
    cleaned = cleaned.replace("www.", "")
    cleaned = cleaned.split(".", 1)[0]
    cleaned = re.sub(r"[^A-Za-z0-9&' -]", "", cleaned).strip()
    if not cleaned:
        cleaned = "Unknown"
    return cleaned[:1].upper() + cleaned[1:].lower()


def next_source_list_path(base_dir):
    index = 1
    while True:
        path = os.path.join(base_dir, f"{EXPORT_PREFIX}{index}.docx")
        if not os.path.exists(path):
            return path
        index += 1


def export_priority_source_list():
    export_file = word_var.get().strip()
    if not export_file:
        messagebox.showwarning("Missing Word File", "Please select a Word file first.")
        return
    try:
        export_doc = Document(export_file)
    except Exception as e:
        messagebox.showwarning("Open Failed", f"Could not open the selected Word file.\n\n{e}")
        return

    output_doc = Document()
    normal_style = output_doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(22)

    entries = []
    for table in export_doc.tables:
        for row in table.rows:
            if len(row.cells) < 4:
                continue
            number = row.cells[1].text.strip()
            title = row.cells[2].text.strip()
            url = row.cells[3].text.strip()
            if not title:
                continue
            parsed = urlparse(url) if url else None
            source = parsed.netloc.replace("www.", "") if parsed else "Unknown"
            entries.append((number, title, source_display_name(source)))

    if not entries:
        messagebox.showinfo("No News Found", "No valid news rows were found in the selected file.")
        return

    export_path = next_source_list_path(os.path.dirname(export_file))
    for number, title, source_name in entries:
        paragraph = output_doc.add_paragraph()
        run_number = paragraph.add_run(f"#{number}")
        run_number.bold = True
        run_number.font.name = "Times New Roman"
        run_number.font.size = Pt(22)
        run_middle = paragraph.add_run(f"   {title} ")
        run_middle.font.name = "Times New Roman"
        run_middle.font.size = Pt(22)
        run_source = paragraph.add_run(f"({source_name})")
        run_source.bold = True
        run_source.font.name = "Times New Roman"
        run_source.font.size = Pt(22)
    output_doc.save(export_path)
    messagebox.showinfo("Export Complete", f"Source list exported to:\n{export_path}")


def browse_word():
    initial_file = word_var.get().strip()
    initial_dir = os.path.dirname(initial_file) if initial_file else BASE_DIR
    if not os.path.isdir(initial_dir):
        initial_dir = BASE_DIR
    path = filedialog.askopenfilename(initialdir=initial_dir, filetypes=[("Word Documents", "*.docx")])
    if path:
        word_var.set(path)
        if not output_var.get().strip():
            output_var.set(os.path.dirname(path))


def browse_output():
    initial_dir = output_var.get().strip() or os.path.dirname(word_var.get().strip()) or BASE_DIR
    if not os.path.isdir(initial_dir):
        initial_dir = BASE_DIR
    path = filedialog.askdirectory(initialdir=initial_dir, title="Choose PNG render folder")
    if path:
        output_var.set(path)


def prepare_common_state():
    global file_path, output_path, highlight_bold_text, highlight_color, highlight_opacity
    file_path = word_var.get().strip()
    output_path = output_var.get().strip()
    highlight_bold_text = True
    highlight_color = highlight_color_var.get().strip() or DEFAULT_HIGHLIGHT_COLOR
    try:
        highlight_opacity = max(0, min(100, int(highlight_opacity_var.get().strip() or DEFAULT_HIGHLIGHT_OPACITY)))
    except ValueError:
        messagebox.showwarning("Invalid Opacity", "Highlight opacity must be a number from 0 to 100.")
        return False
    if not file_path:
        messagebox.showwarning("Missing Word File", "Please select a Word file.")
        return False
    if not output_path:
        output_path = os.path.dirname(file_path)
        output_var.set(output_path)
    if not os.path.isdir(output_path):
        messagebox.showwarning("Missing Render Folder", "Please choose a render folder for the PNG files.")
        return False
    save_settings(file_path, output_path, highlight_color, highlight_opacity)
    return True


def run_app():
    global action_mode
    if not prepare_common_state():
        return
    action_mode = "render"
    root.quit()


def run_check_only():
    global action_mode
    if not prepare_common_state():
        return
    action_mode = "check"
    root.quit()

settings = load_settings()

root = Tk()
root.title(f"News Image Generator v{APP_VERSION}")
root.geometry("760x380")
root.resizable(False, False)

word_var = StringVar(value=settings.get("last_word_file", ""))
output_var = StringVar(value=settings.get("last_output_folder", ""))
highlight_color_var = StringVar(value=settings.get(HIGHLIGHT_COLOR_SETTINGS_KEY, DEFAULT_HIGHLIGHT_COLOR))
highlight_opacity_var = StringVar(value=str(settings.get(HIGHLIGHT_OPACITY_SETTINGS_KEY, DEFAULT_HIGHLIGHT_OPACITY)))

Label(root, text="PrintNews", font=("Georgia", 18, "bold")).place(x=24, y=18)
Button(root, text="Sync Logos", width=14, command=start_logo_sync).place(x=220, y=18)
Button(root, text="Set GitHub Token", width=16, command=set_github_token).place(x=350, y=18)
update_button = Button(root, text="Checking update...", width=22, command=check_updates_clicked)
update_button.place(x=500, y=18)

Label(root, text="Word File", font=("Georgia", 10, "bold")).place(x=24, y=84)
Entry(root, textvariable=word_var, width=58, relief="solid", bd=1).place(x=24, y=108)
Button(root, text="Browse", command=browse_word).place(x=465, y=104)

Label(root, text="Render Folder", font=("Georgia", 10, "bold")).place(x=24, y=148)
Entry(root, textvariable=output_var, width=58, relief="solid", bd=1).place(x=24, y=172)
Button(root, text="Browse", command=browse_output).place(x=465, y=168)

Label(root, text="Note: Bold your headline text to apply hilight after checking.", font=("Georgia", 10, "bold")).place(x=24, y=228)
Label(root, text="Highlight Color").place(x=24, y=264)
Entry(root, textvariable=highlight_color_var, width=14, relief="solid", bd=1).place(x=120, y=264)
Button(root, text="Choose", command=choose_highlight_color).place(x=238, y=260)
Label(root, text="Opacity %").place(x=330, y=264)
Entry(root, textvariable=highlight_opacity_var, width=8, relief="solid", bd=1).place(x=400, y=264)

Button(root, text="Check News Against Link", width=22, command=run_check_only).place(x=70, y=325)
Button(root, text="Export Priority Source List", width=24, command=export_priority_source_list).place(x=278, y=325)
Button(root, text="Render News", width=22, command=run_app).place(x=520, y=325)
root.after(500, start_update_button_check)

root.mainloop()

if not file_path:
    print("No Word file selected. Exiting.")
    sys.exit()

INPUT_FOLDER = os.path.dirname(file_path)
OUTPUT_FOLDER = output_path

# ---------------- Helpers ----------------
def resize_logo(img):
    ratio = LOGO_HEIGHT / img.height
    new_width = int(img.width * ratio)
    return img.resize((new_width, LOGO_HEIGHT), Image.LANCZOS)

def text_width(font, text):
    return font.getbbox(text)[2]

def wrap_headline(headline, main_font, sub_font, max_width):
    segments = headline.split("//")
    lines = []
    first_segment = True
    for seg in segments:
        seg = seg.strip()
        if not seg:
            continue
        font = main_font if first_segment else sub_font
        color = "black" if first_segment else SUB_HEAD_COLOR
        words = seg.split()
        line = ""
        for word in words:
            test_line = word if line == "" else line + " " + word
            if text_width(font, test_line) + 2*MARGIN > max_width:
                if line:
                    lines.append((line, font, color))
                line = word
            else:
                line = test_line
        if line:
            lines.append((line, font, color))
        # Add extra gap after each // segment except the first
        if not first_segment:
            lines.append(("", sub_font, SUB_HEAD_COLOR))
        first_segment = False
    return lines

def normalize_date(d):
    if re.fullmatch(r'\d{8}', d):
        return f"{d[:4]}-{d[4:6]}-{d[6:]}"
    return d


def normalize_text(value):
    text = (value or "").strip()
    replacements = {
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2013": "-",
        "\u2014": "-",
        "\u00a0": " ",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return re.sub(r"\s+", " ", text).strip()


def normalize_title(value):
    title = normalize_text(value).lower()
    title = re.sub(r"\s+\|\s+[^|]+$", "", title)
    title = re.sub(r"\s+-\s+[^-]+$", "", title)
    title = re.sub(r"[^\w\s]", "", title)
    return re.sub(r"\s+", " ", title).strip()


def parse_date_to_canonical(text):
    raw = normalize_text(text)
    if not raw:
        return ""
    raw = re.sub(r"\b(\d{1,2})(st|nd|rd|th)\b", r"\1", raw, flags=re.IGNORECASE)
    compact = re.fullmatch(r"(\d{4})(\d{2})(\d{2})", raw)
    if compact:
        return f"{compact.group(1)}-{compact.group(2)}-{compact.group(3)}"
    iso = re.search(r"(\d{4})-(\d{2})-(\d{2})", raw)
    if iso:
        return f"{iso.group(1)}-{iso.group(2)}-{iso.group(3)}"
    slash = re.search(r"\b(\d{1,4})[/-](\d{1,2})[/-](\d{1,4})\b", raw)
    if slash:
        a, b, c = int(slash.group(1)), int(slash.group(2)), int(slash.group(3))
        candidates = [(a, b, c)] if a > 999 else []
        if c > 999:
            candidates += [(c, a, b), (c, b, a)]
        for year, month, day in candidates:
            try:
                return datetime(year, month, day).strftime("%Y-%m-%d")
            except ValueError:
                continue
    spelled = re.search(r"\b([A-Za-z]+)\s+(\d{1,2}),?\s+(\d{4})\b", raw)
    if spelled:
        month = MONTH_MAP.get(spelled.group(1).lower())
        if month:
            try:
                return datetime(int(spelled.group(3)), month, int(spelled.group(2))).strftime("%Y-%m-%d")
            except ValueError:
                pass
    spoken = re.search(r"\b(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})\b", raw)
    if spoken:
        month = MONTH_MAP.get(spoken.group(2).lower())
        if month:
            try:
                return datetime(int(spoken.group(3)), month, int(spoken.group(1))).strftime("%Y-%m-%d")
            except ValueError:
                pass
    return ""


def format_date_like_sheet(canonical_date, sheet_date_raw):
    if not canonical_date:
        return ""
    try:
        dt = datetime.strptime(canonical_date, "%Y-%m-%d")
    except ValueError:
        return canonical_date
    sheet_raw = (sheet_date_raw or "").strip()
    if re.fullmatch(r"\d{8}", sheet_raw):
        return dt.strftime("%Y%m%d")
    if re.fullmatch(r"\d{4}-\d{2}-\d{2}", sheet_raw):
        return dt.strftime("%Y-%m-%d")
    if re.fullmatch(r"\d{4}/\d{2}/\d{2}", sheet_raw):
        return dt.strftime("%Y/%m/%d")
    return canonical_date


def title_exact_match(sheet_title, web_title):
    return bool(normalize_title(sheet_title) and normalize_title(sheet_title) == normalize_title(web_title))


def date_exact_match(sheet_date, web_date):
    return bool(parse_date_to_canonical(sheet_date) and parse_date_to_canonical(sheet_date) == parse_date_to_canonical(web_date))

def base_logo_name(name):
    cleaned = (name or "Unknown").strip()
    return cleaned.split(".", 1)[0] if "." in cleaned else cleaned


def missing_logo_note_name(name):
    cleaned = (name or "Unknown").strip()
    parts = cleaned.split(".") if cleaned else []
    return ".".join(parts[:-1]) if len(parts) > 1 else cleaned


def preferred_logo_name(name):
    cleaned = (name or "Unknown").strip()
    preferred = missing_logo_note_name(cleaned)
    return preferred or cleaned


def logo_name_candidates(name):
    names = []
    cleaned = (name or "").strip()
    base_name = base_logo_name(cleaned)
    note_name = missing_logo_note_name(cleaned)
    parts = cleaned.split(".") if cleaned else []
    country_name = base_name + "-" + parts[-1] if len(parts) > 2 else ""
    for candidate in (cleaned, note_name, base_name, country_name, base_name + ".com"):
        if candidate:
            names.append(candidate)
    return list(dict.fromkeys(names))


def missing_logo_name(source):
    return missing_logo_note_name(source)


def missing_logo_search_name(source):
    return (source or "Unknown").strip()


def find_logo_path(name):
    if not os.path.isdir(LOGO_FOLDER):
        return None

    preferred_name = preferred_logo_name(name)
    expected = {
        f"{candidate}{ext}".lower()
        for candidate in logo_name_candidates(name)
        for ext in SUPPORTED_LOGO_EXTENSIONS
    }
    for filename in os.listdir(LOGO_FOLDER):
        if filename.lower() in expected:
            current_path = os.path.join(LOGO_FOLDER, filename)
            current_base, current_ext = os.path.splitext(filename)
            if preferred_name and current_base != preferred_name:
                preferred_filename = f"{preferred_name}{current_ext}"
                preferred_path = os.path.join(LOGO_FOLDER, preferred_filename)
                if not os.path.exists(preferred_path):
                    try:
                        os.replace(current_path, preferred_path)
                        print(f"Renamed logo: {filename} -> {preferred_filename}")
                        return preferred_path
                    except Exception:
                        pass
            return current_path
    return None


def open_logo(path):
    with Image.open(path) as img:
        return resize_logo(img.convert("RGBA"))


def get_logo(source, url=None):
    path = find_logo_path(source)
    if path:
        return open_logo(path), False

    custom_path = find_logo_path(CUSTOM_LOGO_NAME)
    if custom_path:
        return open_logo(custom_path), True

    img = Image.new("RGBA", (LOGO_HEIGHT, LOGO_HEIGHT), (255,255,255,255))
    return img, True

# ---------------- Fonts ----------------
HEADLINE_FONT_FILES = [
    "arialbd.ttf",
    "ARIALNB.TTF",
    "impact.ttf",
    "bahnschrift.ttf",
    "calibrib.ttf",
    "cambriab.ttf",
    "Candarab.ttf",
    "corbelb.ttf",
    "georgiab.ttf",
    "segoeuib.ttf",
    "tahomabd.ttf",
    "timesbd.ttf",
    "trebucbd.ttf",
    "verdanab.ttf",
    "SourceSansPro-Bold.otf",
    "SourceSansPro-Semibold.otf",
    "malgunbd.ttf",
    "msjhbd.ttc",
    "msyhbd.ttc",
    "comicbd.ttf",
]


def load_font(font_file, size):
    font_path = os.path.join("C:/Windows/Fonts", font_file)
    try:
        return ImageFont.truetype(font_path, size)
    except Exception:
        try:
            return ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", size)
        except Exception:
            return ImageFont.load_default()


headline_font_pairs = [
    (load_font(font_file, FONT_SIZE_HEADLINE), load_font(font_file, int(FONT_SIZE_HEADLINE * 0.8)))
    for font_file in HEADLINE_FONT_FILES
]


def headline_fonts_for_index(index):
    return headline_font_pairs[index % len(headline_font_pairs)]


def headline_fonts_for_source(source, source_font_indices, next_index):
    source_key = (source or "Unknown").strip().lower()
    if source_key not in source_font_indices:
        source_font_indices[source_key] = next_index
        next_index += 1
    return headline_fonts_for_index(source_font_indices[source_key]), next_index


font_head, font_sub_head = headline_fonts_for_index(0)
font_date = load_font("arial.ttf", FONT_SIZE_DATE)
font_source = load_font("arial.ttf", FONT_SIZE_SOURCE)


def open_missing_logo_searches(missing_search_sources):
    for source in sorted(missing_search_sources):
        query = quote_plus(f"{source} logo")
        webbrowser.open_new_tab(f"https://www.google.com/search?tbm=isch&q={query}")


def build_highlight_fill(color_hex, opacity_percent):
    color_hex = (color_hex or DEFAULT_HIGHLIGHT_COLOR).strip()
    if not re.fullmatch(r"#[0-9a-fA-F]{6}", color_hex):
        color_hex = DEFAULT_HIGHLIGHT_COLOR
    opacity = max(0, min(100, int(opacity_percent)))
    alpha = int(255 * (opacity / 100.0))
    return tuple(int(color_hex[i:i + 2], 16) for i in (1, 3, 5)) + (alpha,)


def extract_row_runs(cell):
    runs = []
    for paragraph_index, paragraph in enumerate(cell.paragraphs):
        if paragraph_index > 0:
            runs.append({"text": "\n", "bold": False})
        for run in paragraph.runs:
            if run.text:
                runs.append({"text": run.text, "bold": bool(run.bold)})
    if not runs:
        text = cell.text.strip()
        if text:
            runs.append({"text": text, "bold": False})
    return runs


def flatten_segments(segments):
    flat = []
    for segment in segments:
        pieces = segment["text"].split("\n")
        for index, piece in enumerate(pieces):
            if piece:
                flat.append({"text": piece, "bold": segment["bold"]})
            if index < len(pieces) - 1:
                flat.append({"text": "\n", "bold": False})
    return flat


def build_headline_segments(cell):
    runs = extract_row_runs(cell)
    if not runs:
        return [{"text": "", "bold": False, "font": font_head, "color": "black"}]

    full_text = "".join(run["text"] for run in runs)
    parts = full_text.split("//")
    segments = []
    cursor = 0
    for part_index, part in enumerate(parts):
        part_len = len(part)
        bucket = []
        remaining = part_len
        while remaining > 0 and cursor < len(runs):
            current = runs[cursor]
            current_text = current["text"]
            take = min(len(current_text), remaining)
            bucket.append({"text": current_text[:take], "bold": current["bold"]})
            if take == len(current_text):
                cursor += 1
            else:
                runs[cursor] = {"text": current_text[take:], "bold": current["bold"]}
            remaining -= take
        if bucket:
            font = font_head if part_index == 0 else font_sub_head
            color = "black" if part_index == 0 else SUB_HEAD_COLOR
            for entry in flatten_segments(bucket):
                entry["font"] = font
                entry["color"] = color
                segments.append(entry)
        if part_index < len(parts) - 1:
            if cursor < len(runs) and runs[cursor]["text"].startswith("//"):
                if runs[cursor]["text"] == "//":
                    cursor += 1
                else:
                    runs[cursor] = {"text": runs[cursor]["text"][2:], "bold": runs[cursor]["bold"]}
            segments.append({"text": "\n", "bold": False, "font": font_sub_head, "color": SUB_HEAD_COLOR, "segment_break": True})
    return segments or [{"text": cell.text.strip(), "bold": False, "font": font_head, "color": "black"}]


def wrap_styled_segments(segments, max_width):
    lines = []
    current_line = []
    current_width = 0
    pending_gap = False

    def flush_line():
        nonlocal current_line, current_width, pending_gap
        if current_line:
            lines.append({"parts": current_line, "segment_gap_before": pending_gap})
            current_line = []
            current_width = 0
            pending_gap = False

    for segment in segments:
        if segment.get("segment_break"):
            flush_line()
            pending_gap = True
            continue
        if segment["text"] == "\n":
            flush_line()
            continue
        pieces = re.findall(r"\S+\s*|\s+", segment["text"])
        for piece in pieces:
            piece_width = text_width(segment["font"], piece)
            if current_line and current_width + piece_width + 2 * MARGIN > max_width and piece.strip():
                flush_line()
            current_line.append({
                "text": piece,
                "font": segment["font"],
                "color": segment["color"],
                "bold": segment["bold"],
            })
            current_width += piece_width
    flush_line()
    return lines or [{"parts": [{"text": "", "font": font_head, "color": "black", "bold": False}], "segment_gap_before": False}]


def replace_cell_text(cell, new_text):
    cell.text = new_text


def strip_bold_from_cell(cell):
    return


def normalize_runs_for_storage(runs):
    normalized = []
    for run in runs or []:
        text = run.get("text", "")
        if not text:
            continue
        bold = bool(run.get("bold"))
        if normalized and normalized[-1]["bold"] == bold:
            normalized[-1]["text"] += text
        else:
            normalized.append({"text": text, "bold": bold})
    return normalized


def write_runs_to_cell(cell, runs):
    runs = normalize_runs_for_storage(runs)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    if not runs:
        return

    for run_data in runs:
        pieces = run_data["text"].split("\n")
        for piece_index, piece in enumerate(pieces):
            if piece:
                new_run = paragraph.add_run(piece)
                new_run.bold = run_data["bold"]
            if piece_index < len(pieces) - 1:
                paragraph = cell.add_paragraph()


def text_widget_runs(text_widget):
    raw_text = text_widget.get("1.0", "end-1c")
    if not raw_text:
        return []

    def index_to_offset(index):
        return len(text_widget.get("1.0", index))

    bold_spans = []
    ranges = text_widget.tag_ranges("bold")
    for index in range(0, len(ranges), 2):
        start = index_to_offset(ranges[index])
        end = index_to_offset(ranges[index + 1])
        if end > start:
            bold_spans.append((start, end))

    runs = []
    buffer = []
    current_bold = None
    for char_index, char in enumerate(raw_text):
        is_bold = any(start <= char_index < end for start, end in bold_spans)
        if current_bold is None:
            current_bold = is_bold
        if is_bold != current_bold:
            runs.append({"text": "".join(buffer), "bold": current_bold})
            buffer = [char]
            current_bold = is_bold
        else:
            buffer.append(char)
    if buffer:
        runs.append({"text": "".join(buffer), "bold": bool(current_bold)})
    return normalize_runs_for_storage(runs)


def fill_text_widget_with_runs(text_widget, runs):
    text_widget.delete("1.0", "end")
    for run in runs or []:
        start_index = text_widget.index("end-1c")
        text_widget.insert("end", run.get("text", ""))
        end_index = text_widget.index("end-1c")
        if run.get("bold") and text_widget.compare(end_index, ">", start_index):
            text_widget.tag_add("bold", start_index, end_index)


def toggle_text_widget_bold(text_widget):
    try:
        selection_start = text_widget.index("sel.first")
        selection_end = text_widget.index("sel.last")
    except Exception:
        return "break"

    if text_widget.tag_nextrange("bold", selection_start, selection_end):
        text_widget.tag_remove("bold", selection_start, selection_end)
    else:
        text_widget.tag_add("bold", selection_start, selection_end)
    return "break"


def undo_text_widget(text_widget):
    try:
        text_widget.edit_undo()
    except Exception:
        pass
    return "break"


def redo_text_widget(text_widget):
    try:
        text_widget.edit_redo()
    except Exception:
        pass
    return "break"


def append_news_to_sheet(document, item):
    target_table = item.get("table")
    if target_table is None:
        return
    new_row = target_table.add_row()
    if len(new_row.cells) >= 4:
        replace_cell_text(new_row.cells[0], item.get("resolved_date", item["date_raw"]))
        replace_cell_text(new_row.cells[1], "")
        write_runs_to_cell(new_row.cells[2], item.get("resolved_headline_runs", item.get("headline_runs", [])))
        replace_cell_text(new_row.cells[3], item.get("resolved_url", item["url"]))


def collect_dialog_item_values(item, title_box, date_var, link_var):
    resolved_runs = text_widget_runs(title_box)
    resolved_headline = "".join(run["text"] for run in resolved_runs).strip() or item["headline"]
    resolved_date = date_var.get().strip() or item["date_raw"]
    resolved_url = link_var.get().strip() or item["url"]
    parsed = urlparse(resolved_url) if resolved_url else None
    source = parsed.netloc.replace("www.", "") if parsed else "Unknown"
    return {
        "resolved_headline_runs": resolved_runs,
        "resolved_headline": resolved_headline,
        "resolved_date": resolved_date,
        "resolved_url": resolved_url,
        "source": source,
    }


def apply_table_layout(document):
    for table in document.tables:
        table.autofit = False
        for row in table.rows:
            if len(row.cells) >= 4:
                row.cells[0].width = Inches(0.75)
                row.cells[1].width = Inches(0.55)
                row.cells[2].width = Inches(6.2)
                row.cells[3].width = Inches(1.1)


def run_manual_cross_check(rows, document):
    for index, item in enumerate(rows, start=1):
        url = item["url"].strip()
        if url:
            try:
                webbrowser.open_new_tab(url)
            except Exception:
                pass
        dialog = Toplevel(root)
        dialog.title(f"Cross Check News {index}")
        dialog.geometry("860x470")
        dialog.resizable(False, False)
        date_var = StringVar(value=item["date_raw"])
        link_var = StringVar(value=item["url"])
        accepted = {"value": False}
        Label(dialog, text="Edit headline, date, and link - bold text for highlight", font=("Arial", 16, "bold")).place(x=24, y=20)
        Label(dialog, text=f"News {index} of {len(rows)}", font=("Arial", 10)).place(x=24, y=58)
        open_button = Button(dialog, text="Open Link", width=14, command=lambda link=url: webbrowser.open_new_tab(link) if link else None)
        open_button.place(x=710, y=20)
        attach_tooltip(open_button, url)

        Label(dialog, text="News Title", font=("Arial", 11, "bold")).place(x=24, y=100)
        title_box = Text(
            dialog,
            width=98,
            height=8,
            relief="solid",
            bd=1,
            wrap="word",
            font=("Times New Roman", 13),
            undo=True,
            autoseparators=True,
            maxundo=-1,
        )
        title_box.place(x=24, y=126)
        title_box.tag_configure("bold", font=("Times New Roman", 13, "bold"))
        fill_text_widget_with_runs(title_box, item.get("headline_runs", []))
        title_box.edit_reset()
        title_box.edit_separator()

        Label(dialog, text="News Date", font=("Arial", 11, "bold")).place(x=24, y=284)
        Entry(dialog, textvariable=date_var, width=28, relief="solid", bd=1, font=("Times New Roman", 13)).place(x=24, y=334)

        Label(dialog, text="Source Link", font=("Arial", 11, "bold")).place(x=24, y=372)
        Entry(dialog, textvariable=link_var, width=96, relief="solid", bd=1, font=("Times New Roman", 12)).place(x=24, y=402)
        title_box.bind("<Control-b>", lambda event, box=title_box: toggle_text_widget_bold(box))
        title_box.bind("<Control-B>", lambda event, box=title_box: toggle_text_widget_bold(box))
        title_box.bind("<Control-z>", lambda event, box=title_box: undo_text_widget(box))
        title_box.bind("<Control-Z>", lambda event, box=title_box: undo_text_widget(box))
        title_box.bind("<Control-y>", lambda event, box=title_box: redo_text_widget(box))
        title_box.bind("<Control-Y>", lambda event, box=title_box: redo_text_widget(box))

        def flash_add_button():
            add_to_sheet_button.config(text="Added")
            dialog.after(700, lambda: add_to_sheet_button.config(text="Add This News To Sheet"))

        def add_this_news_now():
            values = collect_dialog_item_values(item, title_box, date_var, link_var)
            item.update(values)
            append_news_to_sheet(document, item)
            apply_table_layout(document)
            document.save(file_path)
            flash_add_button()

        def confirm():
            item.update(collect_dialog_item_values(item, title_box, date_var, link_var))
            item["append_to_sheet"] = False
            item["skipped"] = False
            accepted["value"] = True
            dialog.destroy()

        def skip_news():
            item["append_to_sheet"] = False
            item["resolved_headline"] = ""
            item["resolved_date"] = ""
            item["resolved_url"] = ""
            item["source"] = "Unknown"
            item["resolved_headline_runs"] = []
            item["skipped"] = True
            accepted["value"] = True
            dialog.destroy()

        def cancel():
            if messagebox.askyesno("Cancel Cross Check", "Cancel the render?", parent=dialog):
                dialog.destroy()

        Button(dialog, text="Bold Selected Text", width=18, command=lambda box=title_box: toggle_text_widget_bold(box)).place(x=250, y=96)
        Button(dialog, text="Undo", width=10, command=lambda box=title_box: undo_text_widget(box)).place(x=430, y=96)
        Button(dialog, text="Redo", width=10, command=lambda box=title_box: redo_text_widget(box)).place(x=520, y=96)
        add_to_sheet_button = Button(dialog, text="Add This News To Sheet", width=24, command=add_this_news_now)
        add_to_sheet_button.place(x=610, y=96)
        Button(dialog, text="Remove This News From Sheet", width=24, command=skip_news).place(x=315, y=334)
        Button(dialog, text="OK", width=16, command=confirm).place(x=540, y=334)
        Button(dialog, text="Cancel", width=16, command=cancel).place(x=700, y=334)
        dialog.grab_set()
        dialog.wait_window()
        if not accepted["value"]:
            return False

    for item in rows:
        if item.get("skipped"):
            replace_cell_text(item["date_cell"], "")
            replace_cell_text(item["number_cell"], "")
            replace_cell_text(item["headline_cell"], "")
            replace_cell_text(item["link_cell"], "")
        else:
            write_runs_to_cell(item["headline_cell"], item.get("resolved_headline_runs", item.get("headline_runs", [])))
            replace_cell_text(item["date_cell"], item.get("resolved_date", item["date_raw"]))
            replace_cell_text(item["link_cell"], item.get("resolved_url", item["url"]))
            strip_bold_from_cell(item["headline_cell"])
        if item.get("append_to_sheet"):
            append_news_to_sheet(document, item)
    apply_table_layout(document)
    document.save(file_path)
    return True

doc = Document(file_path)
rows = []
for table in doc.tables:
    for row in table.rows:
        if len(row.cells) < 4:
            continue
        date_raw = row.cells[0].text.strip()
        number = row.cells[1].text.strip()
        headline = row.cells[2].text.strip()
        url = row.cells[3].text.strip()
        parsed = urlparse(url) if url else None
        if not headline:
            continue
        rows.append({
            "date_raw": date_raw,
            "number": number,
            "headline": headline,
            "headline_runs": extract_row_runs(row.cells[2]),
            "url": url,
            "source": parsed.netloc.replace("www.", "") if parsed else "Unknown",
            "table": table,
            "date_cell": row.cells[0],
            "number_cell": row.cells[1],
            "headline_cell": row.cells[2],
            "link_cell": row.cells[3],
            "resolved_headline": headline,
            "resolved_headline_runs": extract_row_runs(row.cells[2]),
            "resolved_date": date_raw,
            "resolved_url": url,
            "append_to_sheet": False,
            "skipped": False,
        })

if action_mode == "check":
    if not run_manual_cross_check(rows, doc):
        print("Check cancelled.")
        sys.exit()
    print("Check is done.")
    try:
        done_root = Tk()
        done_root.withdraw()
        messagebox.showinfo("Check Complete", "Check is done. The Word file was updated.")
        done_root.destroy()
    except Exception:
        pass
    sys.exit()

missing_sources = set()
missing_search_sources = set()
headline_index = 0
source_font_indices = {}

for item in rows:
    try:
        date_raw = item.get("resolved_date", item["date_raw"])
        number = item["number"]
        headline = item.get("resolved_headline", item["headline"])
        url = item["url"]
        source = item["source"]
        headline_segments = build_headline_segments(item["headline_cell"])

        (font_head, font_sub_head), headline_index = headline_fonts_for_source(
            source,
            source_font_indices,
            headline_index,
        )

        words = headline.split()[:MAX_FILENAME_WORDS]
        name_base = re.sub(r'[\/:*?"<>|]', '', f"{number} {' '.join(words)}")[:120]
        name = name_base

        logo, used_fallback = get_logo(source, url)
        if used_fallback:
            missing_sources.add(missing_logo_name(source))
            missing_search_sources.add(missing_logo_search_name(source))

        save_path = os.path.join(OUTPUT_FOLDER, f"{name}.png")
        date = normalize_date(date_raw)

        width = MIN_WIDTH
        for _ in range(10):
            lines = wrap_styled_segments(headline_segments, width)
            if len(lines) <= 2 or width >= MAX_WIDTH:
                break
            width += 50

        text_h = 0
        for line in lines:
            fonts_in_line = [part["font"] for part in line["parts"] if part["text"]]
            line_height = max((font.size for font in fonts_in_line), default=font_head.size)
            text_h += line_height + LINE_SPACING + (GAP_BETWEEN_SEGMENTS if line.get("segment_gap_before") else 0)
        logo_h = logo.height
        height = PADDING_TOP + logo_h + 20 + text_h + PADDING_BOTTOM
        max_line_width = 0
        for line in lines:
            line_width = sum(text_width(part["font"], part["text"]) for part in line["parts"])
            max_line_width = max(max_line_width, line_width)
        date_source_width = text_width(font_date, date) + 10 + text_width(font_source, f" | {source}")
        logo_plus_spacing = logo.width + 20
        final_w = max(max_line_width + MARGIN, logo_plus_spacing + date_source_width + MARGIN)
        final_w += RIGHT_MARGIN

        img = Image.new("RGB", (final_w, height), "white")
        draw = ImageDraw.Draw(img)
        y = PADDING_TOP
        img.paste(logo, (MARGIN, y), logo)
        logo_bottom = y + logo.height
        dx = MARGIN + logo.width + 20
        dy = y + (logo.height - FONT_SIZE_DATE)
        draw.text((dx, dy), date, font=font_date, fill="black")
        dw = text_width(font_date, date)
        draw.text((dx + dw + 10, dy), f" | {source}", font=font_source, fill="black")
        y_text = logo_bottom + 20
        highlight_fill = build_highlight_fill(highlight_color, highlight_opacity)
        for line in lines:
            if line.get("segment_gap_before"):
                y_text += GAP_BETWEEN_SEGMENTS
            x_text = MARGIN
            line_fonts = [part["font"] for part in line["parts"] if part["text"]]
            line_height = max((font.size for font in line_fonts), default=font_head.size)
            for part in line["parts"]:
                if not part["text"]:
                    continue
                part_width = text_width(part["font"], part["text"])
                if highlight_bold_text and part.get("bold"):
                    overlay = Image.new("RGBA", img.size, (255, 255, 255, 0))
                    overlay_draw = ImageDraw.Draw(overlay)
                    overlay_draw.rounded_rectangle(
                        (
                            x_text - HIGHLIGHT_PADDING,
                            y_text - HIGHLIGHT_PADDING,
                            x_text + part_width + HIGHLIGHT_PADDING,
                            y_text + part["font"].size + HIGHLIGHT_PADDING,
                        ),
                        radius=4,
                        fill=highlight_fill,
                    )
                    img = Image.alpha_composite(img.convert("RGBA"), overlay).convert("RGB")
                    draw = ImageDraw.Draw(img)
                draw.text((x_text, y_text), part["text"], font=part["font"], fill=part["color"])
                x_text += part_width
            y_text += line_height + LINE_SPACING
        img.save(save_path)
        print("Saved:", save_path)
    except Exception as e:
        print("Error:", e)

# ---------------- Save missing logos ----------------
with open(MISSING_LOG_FILE, "w", encoding="utf-8") as f:
    for s in sorted(missing_sources):
        f.write(s + "\n")
if missing_sources:
    print(f"Missing logos saved to: {MISSING_LOG_FILE}")
    open_missing_logo_searches(missing_search_sources)

print("Render is done.")
try:
    messagebox.showinfo("Render Complete", "Render is done.")
except Exception:
    pass
















